import sys
from PyQt5.QtWidgets import QMainWindow, QWidget, QPushButton, QApplication, QGridLayout, QLabel, QLineEdit
from PyQt5.QtGui import QIcon, QPixmap
from PyQt5.QtCore import QSize
from math import pi

import os

from docx import Document
from docx.shared import Inches


class FirstWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.WindowInfo = None
        self.build()

    def build(self):

        self.L_label = QLabel("P - образ", self)
        self.L_label.resize(80, 20)
        self.L_label.move(100, 20)

        self.L_label = QLabel("M - образ", self)
        self.L_label.resize(80, 20)
        self.L_label.move(320, 20)

        self.L_label = QLabel("W - образ", self)
        self.L_label.resize(80, 20)
        self.L_label.move(540, 20)

        self.P_btn = QPushButton(self)
        self.P_btn.resize(200, 180)
        self.P_btn.move(20, 50)
        self.P_btn.setIcon(QIcon('img/P_img.png'))
        self.P_btn.setIconSize(QSize(190, 190))
        self.P_btn.clicked.connect(self.show_windiw2_mod_P)

        self.M_btn = QPushButton(self)
        self.M_btn.resize(200, 180)
        self.M_btn.move(240, 50)
        self.M_btn.setIcon(QIcon('img/M_img.png'))
        self.M_btn.setIconSize(QSize(180, 180))
        self.M_btn.clicked.connect(self.show_windiw2_mod_M)

        self.W_btn = QPushButton(self)
        self.W_btn.resize(200, 180)
        self.W_btn.move(460, 50)
        self.W_btn.setIcon(QIcon('img/W_img.png'))
        self.W_btn.setIconSize(QSize(180, 180))
        self.W_btn.clicked.connect(self.show_windiw2_mod_W)

        self.setGeometry(300, 200, 680, 250)
        self.setWindowTitle('GCI')
        self.setWindowIcon(QIcon('img/logo.jpg'))
        self.show()

    def show_windiw2_mod_P(self):
        if not self.WindowInfo:
            self.WindowInfo = WindowInfo(mode=5)
        elif self.WindowInfo and not self.WindowInfo.isVisible():
            self.WindowInfo = WindowInfo(mode=5)
        self.WindowInfo.setWindowTitle("PWindow")
        self.WindowInfo.show()

    def show_windiw2_mod_M(self):
        if not self.WindowInfo:
            self.WindowInfo = WindowInfo(mode=6)
        elif self.WindowInfo and not self.WindowInfo.isVisible():
            self.WindowInfo = WindowInfo(mode=6)
        self.WindowInfo.setWindowTitle("MWindow")
        self.WindowInfo.show()

    def show_windiw2_mod_W(self):
        if not self.WindowInfo:
            self.WindowInfo = WindowInfo(mode=8)
        elif self.WindowInfo and not self.WindowInfo.isVisible():
            self.WindowInfo = WindowInfo(mode=8)
        self.WindowInfo.setWindowTitle("WWindow")
        self.WindowInfo.show()


class WindowInfo(QWidget):
    def __init__(self, mode=0):
        super().__init__()
        self.ResultsWindow = None
        self.mode = mode
        self.build()

    def build(self):
        grid = QGridLayout()
        grid.setSpacing(20)

        self.picture = QLabel(self)

        if self.mode == 5:
            pixmap = QPixmap('img/P_drow.jpg')

        elif self.mode == 6:
            pixmap = QPixmap('img/M_drow.jpg')

        else:
            pixmap = QPixmap('img/W_drow.jpg')

        self.picture.setPixmap(pixmap)
        self.resize(pixmap.width(), pixmap.height())

        grid.addWidget(self.picture, 0, 0, 1, 4)

        self.l = []         # список виджетов прямых участков
        self.values_list = []   # список значений прямых участков

        for i in range(2, self.mode):
            self.L = QLabel("L" + str(i - 1), self)
            self.L_inp = QLineEdit("0", self)
            grid.addWidget(self.L, i, 0)
            grid.addWidget(self.L_inp, i, 1)
            self.l.append(self.L_inp)

        self.R = QLabel("R =", self)
        grid.addWidget(self.R, 2, 2)
        self.R_inp = QLineEdit("0", self)
        grid.addWidget(self.R_inp, 2, 3)

        self.D = QLabel("D =", self)
        grid.addWidget(self.D, 3, 2)
        self.D_inp = QLineEdit("0", self)
        grid.addWidget(self.D_inp, 3, 3)

        self.but_L0 = QPushButton('Длина изделия ( L0 )')
        self.but_L0.clicked.connect(self.calc_length)
        grid.addWidget(self.but_L0, 4, 2, 1, 2)

        self.L0 = QLabel("L0 =", self)
        grid.addWidget(self.L0, 5, 2)
        self.L0_inp = QLabel("0", self)
        grid.addWidget(self.L0_inp, 5, 3)

        self.butResult = QPushButton('Длина развёртки')
        self.butResult.clicked.connect(self.calc_results)
        grid.addWidget(self.butResult, 6, 2, 1, 2)

        self.setLayout(grid)

        self.setGeometry(400, 50, 300, 300)

    def calc_length(self):
        self.L0_inp.setText(str(float(self.l[0].text()) + float(self.R_inp.text()) + float(self.D_inp.text())/2))

    def calc_results(self):
        if not self.ResultsWindow:
            self.ResultsWindow = ResultsWindow(self.mode)
        for i in self.l:
            self.values_list.append(float(i.text()))

        if self.mode == 5:                          # угол гиба = 90 градусов
            self.arc_length = pi * float(self.R_inp.text()) / 2
        else:                                       # угол гиба = 180 градусов
            self.arc_length = pi * float(self.R_inp.text())

        self.ResultsWindow.Total_length_inp.setText(str("%.2f" % (sum(self.values_list) + self.arc_length * (self.mode - 3))))

        self.hoarder = 0

        for i in range(len(self.ResultsWindow.bending_points_list)):
            if i == 0:
                self.ResultsWindow.bending_points_list[i].setText(str("%.2f" % (self.hoarder + self.values_list[i])))
                self.hoarder = self.hoarder + self.values_list[i]
                #   print("hoarder = ", self.hoarder)
            else:
                self.ResultsWindow.bending_points_list[i].setText(str("%.2f" % (self.hoarder + self.values_list[i] +
                                                                                self.arc_length)))
                self.hoarder = self.hoarder + self.values_list[i] + self.arc_length
                #   print("hoarder = ", self.hoarder)

        self.ResultsWindow.show()


class ResultsWindow(QWidget):
    def __init__(self, mode=0):
        super().__init__()
        self.mode = mode
        self.build()

    def build(self):
        grid = QGridLayout()
        grid.setSpacing(20)

        self.pic_label = QLabel(self)

        if self.mode == 5:
            pixmap = QPixmap('img/Result_draw_P.jpg')
        elif self.mode == 6:
            pixmap = QPixmap('img/Result_draw_M.jpg')
        else:
            pixmap = QPixmap('img/Result_draw_W.jpg')

        self.pic_label.setPixmap(pixmap)
        self.resize(pixmap.width(), pixmap.height())
        grid.addWidget(self.pic_label, 0, 0, 1, 4)

        self.Total_length = QLabel("Длина развёртки (L0) = ", self)
        self.Total_length_inp = QLineEdit("0", self)
        grid.addWidget(self.Total_length, 2, 0)
        grid.addWidget(self.Total_length_inp, 2, 1)

        self.bending_points_list = []

        for i in range(2, self.mode - 1):
            self.Bending_point = QLabel("Точка гиба №" + str(i - 1) + "( L" + str(i - 1) + " ) = ", self)
            self.Bending_point_inp = QLineEdit("0", self)
            grid.addWidget(self.Bending_point, i, 2)
            grid.addWidget(self.Bending_point_inp, i, 3)
            self.bending_points_list.append(self.Bending_point_inp)

        self.but_1 = QPushButton('Импортировать в .docx', self)
        self.but_1.clicked.connect(self.create_document)
        grid.addWidget(self.but_1, 7, 0, 1, 4)

        self.setLayout(grid)
        self.setGeometry(250, 250, 300, 300)

    def create_document(self):
        document = Document()
        document.add_heading('Длина развёртки и точки гиба', 0)

        if self.mode == 5:
            document.add_picture('img/Result_draw_P.jpg', width=Inches(7.25))
        elif self.mode == 6:
            document.add_picture('img/Result_draw_M.jpg', width=Inches(7.25))
        else:
            document.add_picture('img/Result_draw_W.jpg', width=Inches(7.25))

        for i in self.bending_points_list:
            document.add_paragraph('Точка гиба №' + str(self.bending_points_list.index(i) + 1) + '(L' +
                                   str(self.bending_points_list.index(i) + 1) + ')= ' + i.text() + 'мм')

        document.add_paragraph('Длина развёртки (L0) = ' + self.Total_length_inp.text() + 'мм')
        document.add_page_break()
        document.save('bend_points.docx')
        os.startfile('bend_points.docx')


if __name__ == '__main__':
    app = QApplication(sys.argv)
    ex = FirstWindow()
    sys.exit(app.exec_())
